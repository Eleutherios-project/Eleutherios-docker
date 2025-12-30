#!/usr/bin/env python3
"""
MVP Multi-Format Document to JSONL Converter - PARALLEL VERSION
Supports: PDF, TXT, EPUB, DOCX
Extends mvp_pdf_to_jsonl_parallel.py to handle multiple formats
"""

import json
import re
from pathlib import Path
from tqdm import tqdm
from multiprocessing import Pool, cpu_count
from functools import partial


# OCR support for scanned PDFs
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

def extract_text_with_ocr(pdf_path: str, dpi: int = 200) -> str:
    """Extract text from scanned PDF using OCR"""
    try:
        text_blocks = []
        images = convert_from_path(pdf_path, dpi=dpi)
        
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img)
            if text.strip():
                text = re.sub(r'\n\s*\n', '\n\n', text)
                text = re.sub(r'[ \t]+', ' ', text)
                text_blocks.append(text)
            
            if (i + 1) % 10 == 0:
                print(f"    OCR progress: {i+1}/{len(images)} pages")
        
        print(f"  OCR complete: {len(text_blocks)} pages with text")
        return '\n\n'.join(text_blocks)
    except Exception as e:
        print(f"  OCR error: {e}")
        return ""

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF with OCR fallback for scanned documents"""
    try:
        import pypdf
        text_blocks = []
        with open(file_path, 'rb') as file:
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text = re.sub(r'\n\s*\n', '\n\n', text)
                    text = re.sub(r'[ \t]+', ' ', text)
                    text_blocks.append(text)
        
        full_text = '\n\n'.join(text_blocks)
        
        # Check text quality - garbage PDFs have mostly non-printable chars
        printable_ratio = sum(1 for c in full_text if c.isprintable() or c.isspace()) / max(len(full_text), 1)
        
        # Trigger OCR if too short OR too much garbage (< 50% printable)
        if (len(full_text.strip()) < 100 or printable_ratio < 0.5) and OCR_AVAILABLE:
            from pathlib import Path
            print(f"  Low quality text ({printable_ratio:.0%} printable), trying OCR: {Path(file_path).name}")
            full_text = extract_text_with_ocr(file_path)
        
        return full_text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text
    except Exception as e:
        print(f"TXT extraction error: {e}")
        return ""


def extract_text_from_epub(file_path: str) -> str:
    """Extract text from EPUB"""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        
        book = epub.read_epub(file_path)
        text_blocks = []
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content()
                soup = BeautifulSoup(content, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                # Clean up
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                if text.strip():
                    text_blocks.append(text)
        
        return '\n\n'.join(text_blocks)
    except ImportError:
        print("EPUB support requires: pip install ebooklib beautifulsoup4")
        return ""
    except Exception as e:
        print(f"EPUB extraction error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_blocks = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_blocks.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_blocks.append(row_text)
        
        text = '\n\n'.join(text_blocks)
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        return text
    except ImportError:
        print("DOCX support requires: pip install python-docx")
        return ""
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_text(file_path: Path) -> tuple[str, str]:
    """
    Extract text from file based on extension
    Returns: (text, source_type)
    """
    suffix = file_path.suffix.lower()
    
    extractors = {
        '.pdf': ('pdf', extract_text_from_pdf),
        '.txt': ('txt', extract_text_from_txt),
        '.epub': ('epub', extract_text_from_epub),
        '.docx': ('docx', extract_text_from_docx),
        '.doc': ('docx', extract_text_from_docx),  # Try same extractor
    }
    
    if suffix in extractors:
        source_type, extractor = extractors[suffix]
        text = extractor(str(file_path))
        return text, source_type
    else:
        return "", "unknown"


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    """Split text into overlapping chunks by words"""
    words = text.split()
    chunks = []
    
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        if len(chunk_words) < 50:  # Skip tiny chunks
            break
        chunks.append(' '.join(chunk_words))
        i += (chunk_size - overlap)
    
    return chunks


def process_single_file(file_path: Path, output_dir: Path, chunk_size: int, overlap: int):
    """Process a single document - designed for parallel execution"""
    try:
        # Extract text
        text, source_type = extract_text(file_path)
        
        if not text:
            return (file_path.name, source_type, 0, "No text extracted")
        
        # Create chunks
        chunks = chunk_text(text, chunk_size, overlap)
        if not chunks:
            return (file_path.name, source_type, 0, "No chunks created")
        
        # Write JSONL
        output_path = output_dir / f"{file_path.stem}.jsonl"
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, chunk in enumerate(chunks):
                entry = {
                    "text": chunk,
                    "metadata": {
                        "source": file_path.name,
                        "source_type": source_type,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "doc_id": f"{file_path.stem}_{i:04d}"
                    }
                }
                f.write(json.dumps(entry) + '\n')
        
        return (file_path.name, source_type, len(chunks), "success")
        
    except Exception as e:
        return (file_path.name, "unknown", 0, f"Error: {str(e)}")


def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Convert documents to JSONL (Parallel)',
        epilog='Supported formats: PDF, TXT, EPUB, DOCX'
    )
    parser.add_argument('input_dir', help='Directory containing documents')
    parser.add_argument('output_dir', help='Output directory for JSONL files')
    parser.add_argument('--chunk-size', type=int, default=1000, help='Words per chunk')
    parser.add_argument('--overlap', type=int, default=200, help='Overlap in words')
    parser.add_argument('--workers', type=int, default=None, 
                        help='Number of parallel workers (default: CPU count - 1)')
    parser.add_argument('--formats', nargs='+', 
                        default=['pdf', 'txt', 'epub', 'docx'],
                        help='File formats to process (default: all)')
    
    args = parser.parse_args()
    
    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find documents (skip Mac resource forks)
    all_files = []
    format_counts = {}
    
    for fmt in args.formats:
        pattern = f"*.{fmt.lower()}"
        files = [f for f in input_dir.glob(pattern) if not f.name.startswith('._')]
        all_files.extend(files)
        if files:
            format_counts[fmt.upper()] = len(files)
    
    if not all_files:
        print(f"❌ No documents found in {input_dir}")
        print(f"   Looking for: {', '.join(args.formats)}")
        return
    
    # Remove duplicates (in case of overlapping patterns)
    all_files = list(set(all_files))
    
    # Determine number of workers
    num_workers = args.workers or max(1, cpu_count() - 1)
    
    print(f"Found {len(all_files)} documents:")
    for fmt, count in sorted(format_counts.items()):
        print(f"  {fmt}: {count} files")
    print(f"Using {num_workers} parallel workers")
    print("=" * 80)
    
    # Create partial function with fixed parameters
    process_func = partial(
        process_single_file, 
        output_dir=output_dir, 
        chunk_size=args.chunk_size, 
        overlap=args.overlap
    )
    
    # Process in parallel with progress bar
    total_chunks = 0
    successful = 0
    failed = 0
    type_stats = {}
    
    with Pool(num_workers) as pool:
        results = list(tqdm(
            pool.imap(process_func, all_files),
            total=len(all_files),
            desc="Converting documents",
            unit="file"
        ))
    
    # Summarize results
    print("\n" + "=" * 80)
    print("RESULTS:")
    print("-" * 80)
    
    for filename, source_type, chunk_count, status in results:
        if status == "success":
            total_chunks += chunk_count
            successful += 1
            type_stats[source_type] = type_stats.get(source_type, 0) + chunk_count
            print(f"✓ {filename} ({source_type}): {chunk_count} chunks")
        else:
            failed += 1
            print(f"✗ {filename}: {status}")
    
    print("=" * 80)
    print(f"✅ Successfully converted: {successful}/{len(all_files)} documents")
    if failed > 0:
        print(f"⚠️  Failed: {failed} documents")
    print(f"\n📊 Chunks by format:")
    for fmt, count in sorted(type_stats.items()):
        print(f"  {fmt.upper()}: {count} chunks")
    print(f"\n📝 Total chunks created: {total_chunks}")
    print(f"📁 Output directory: {output_dir}")
    print("=" * 80)


if __name__ == '__main__':
    main()
